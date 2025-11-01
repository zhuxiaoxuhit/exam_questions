"""
考题生成系统
从输入的xlsx文件中读取鉴定点信息,使用LLM生成考题,输出符合格式要求的xlsx文件
"""

import os
import sys
import io
import re
import json
import logging
from datetime import datetime
from typing import Dict, List, Tuple, Optional
import openpyxl
from openpyxl import Workbook
import xlwt
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
from dotenv import load_dotenv

# Windows下设置UTF-8编码
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# 加载环境变量
load_dotenv()


def setup_logging(log_dir: str = "logs") -> str:
    """
    配置日志系统，同时输出到控制台和文件

    Args:
        log_dir: 日志文件目录

    Returns:
        日志文件路径
    """
    # 创建日志目录
    if not os.path.exists(log_dir):
        os.makedirs(log_dir)

    # 生成日志文件名（带时间戳）
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_filename = f"question_generation_{timestamp}.log"
    log_filepath = os.path.join(log_dir, log_filename)

    # 创建logger
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)

    # 清除已有的handlers
    logger.handlers.clear()

    # 创建文件handler（禁用缓冲）
    file_handler = logging.FileHandler(log_filepath, encoding='utf-8', mode='w')
    file_handler.setLevel(logging.INFO)

    # 创建控制台handler（禁用缓冲）
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    console_handler.flush = sys.stdout.flush  # 强制刷新

    # 创建formatter
    formatter = logging.Formatter('%(message)s')
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)

    # 添加handlers
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)

    return log_filepath


class QuestionGenerator:
    """考题生成器"""

    # 不同等级的出题数量要求（注意：等级仅决定题目数量，不影响题目难度）
    LEVEL_REQUIREMENTS = {
        "一级": {"单选": 3, "判断": 2, "多选": 2},  # 共7题
        "二级": {"单选": 3, "判断": 2, "多选": 2},  # 共7题
        "三级": {"单选": 3, "判断": 2, "多选": 2},  # 共7题
        "四级": {"单选": 4, "判断": 2, "多选": 0},  # 共6题
        "五级": {"单选": 4, "判断": 2, "多选": 0},  # 共6题
    }

    def __init__(self):
        """初始化生成器"""
        api_key = os.getenv("DASHSCOPE_API_KEY")
        if not api_key:
            raise ValueError("请在.env文件中设置DASHSCOPE_API_KEY")

        self.client = OpenAI(
            api_key=api_key,
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
        )

    def extract_level_from_filename(self, filename: str) -> Optional[str]:
        """
        从文件名中提取等级信息

        Args:
            filename: 文件名,例如 "三级3001-3010.xlsx"

        Returns:
            等级字符串,如 "三级",如果未找到则返回None
        """
        # 匹配 "一级", "二级", "三级", "四级", "五级"
        pattern = r'(一级|二级|三级|四级|五级)'
        match = re.search(pattern, filename)

        if match:
            return match.group(1)
        else:
            return None

    def detect_file_format(self, filepath: str) -> str:
        """
        检测xlsx文件的格式类型

        Args:
            filepath: xlsx文件路径

        Returns:
            格式类型: "format1"、"format2" 或 "format3"
        """
        wb = openpyxl.load_workbook(filepath)
        ws = wb.active

        # 读取前3行判断格式
        rows = list(ws.iter_rows(min_row=1, max_row=3, values_only=True))

        # 检查第1行是否包含表头关键字
        if len(rows) >= 1:
            row1 = rows[0]
            row1_str = ' '.join([str(cell) if cell else '' for cell in row1])
            
            # 检查是否为format3（四级文件格式）
            # 特征：第1行包含"鉴定范围"且第6列为"题目序号"
            if '鉴定范围' in row1_str and len(row1) > 5:
                col6_str = str(row1[5]) if row1[5] else ''
                if '题目序号' in col6_str:
                    wb.close()
                    return "format3"
            
            # 如果第1行包含"题目序号"、"鉴定点"、"资料"等关键字，判定为格式2
            if any(keyword in row1_str for keyword in ['题目序号', '鉴定点', '资料']):
                wb.close()
                return "format2"

        # 检查第2行是否包含表头关键字（format2的另一种情况）
        if len(rows) >= 2:
            row2 = rows[1]
            row2_str = ' '.join([str(cell) if cell else '' for cell in row2])
            # 如果第2行包含"题目序号"、"鉴定点"、"资料"等关键字，判定为格式2
            if any(keyword in row2_str for keyword in ['题目序号', '鉴定点', '资料']):
                wb.close()
                return "format2"

        # 检查第1行第1列是否为数字（序号）- format1的特征
        if len(rows) >= 1 and rows[0][0]:
            first_cell = str(rows[0][0])
            if first_cell.isdigit():
                wb.close()
                return "format1"

        wb.close()
        return "format1"  # 默认格式1

    def read_knowledge_points(self, filepath: str) -> List[Dict[str, str]]:
        """
        读取xlsx文件中的鉴定点信息（自动检测格式）

        Args:
            filepath: xlsx文件路径

        Returns:
            鉴定点列表,每个元素包含 {"编号": "...", "名称": "...", "内容": "..."}
        """
        file_format = self.detect_file_format(filepath)
        logging.info(f"  检测到文件格式: {file_format}")

        wb = openpyxl.load_workbook(filepath)
        ws = wb.active

        knowledge_points = []

        if file_format == "format1":
            # 格式1: 第1列=序号, 第2列=编号, 第3列=名称, 第4列=内容
            for row in ws.iter_rows(min_row=1, values_only=True):
                if row[0] is None:  # 跳过空行
                    continue

                # 获取编号并标准化（将中文破折号替换为英文连字符）
                code = str(row[1]) if row[1] else ""
                code = code.replace('—', '-').replace('－', '-')

                point = {
                    "序号": str(row[0]) if row[0] else "",
                    "编号": code,  # 使用标准化后的编号
                    "名称": str(row[2]) if row[2] else "",
                    "内容": str(row[3]) if row[3] else "",
                }
                knowledge_points.append(point)

        elif file_format == "format2":
            # 格式2: 跳过第1行（表头）, 第1列=编号, 第2列=名称, 第3列=内容（索引2）
            for row in ws.iter_rows(min_row=2, values_only=True):
                # 跳过空行（第1列为空）
                if row[0] is None or str(row[0]).strip() == "":
                    continue

                # 获取编号并标准化（将中文破折号替换为英文连字符）
                code = str(row[0]).strip()
                # 替换中文破折号为英文连字符
                code = code.replace('—', '-').replace('－', '-')
                
                # 检查是否为有效的鉴定点编号（格式如 A-B-A-001）
                if not code or len(code) < 5:
                    continue
                
                # 跳过表头行（如果编号包含"题目序号"等关键字）
                if any(keyword in code for keyword in ['题目序号', '鉴定点', '序号']):
                    continue

                point = {
                    "序号": code,  # 第1列就是序号/编号
                    "编号": code,  # 使用标准化后的编号
                    "名称": str(row[1]).strip() if len(row) > 1 and row[1] else "",
                    "内容": str(row[2]).strip() if len(row) > 2 and row[2] else "",  # 第3列是内容
                }
                knowledge_points.append(point)

        else:  # format3
            # 格式3: 四级文件格式，跳过第1行（表头）
            # 第6列（索引5）=编号, 第7列（索引6）=名称, 第9列（索引8）=内容
            for row in ws.iter_rows(min_row=2, values_only=True):
                # 跳过空行（第6列为空）
                if len(row) <= 5 or row[5] is None or str(row[5]).strip() == "":
                    continue

                # 获取编号并标准化（将中文破折号替换为英文连字符）
                code = str(row[5]).strip()
                # 替换中文破折号为英文连字符
                code = code.replace('—', '-').replace('－', '-')
                
                # 检查是否为有效的鉴定点编号（格式如 B-D-A-001）
                if not code or len(code) < 5 or '-' not in code:
                    continue
                
                # 跳过表头行
                if any(keyword in code for keyword in ['题目序号', '鉴定点', '序号']):
                    continue

                point = {
                    "序号": code,  # 使用编号作为序号
                    "编号": code,  # 使用标准化后的编号
                    "名称": str(row[6]).strip() if len(row) > 6 and row[6] else "",
                    "内容": str(row[8]).strip() if len(row) > 8 and row[8] else "",
                }
                knowledge_points.append(point)

        wb.close()
        return knowledge_points


    def generate_all_questions_at_once(self, knowledge_point: Dict[str, str],
                                      level: str) -> Optional[List[Dict]]:
        """
        一次性生成一个鉴定点的所有题目（避免重复）

        Args:
            knowledge_point: 鉴定点信息
            level: 等级（仅用于确定题目数量，不影响难度）

        Returns:
            题目列表，如果生成失败返回None
        """
        # 检查知识点内容是否为空
        content = knowledge_point.get('内容', '').strip()
        if not content:
            logging.error(f"  ❌ 错误：鉴定点内容为空，无法生成题目")
            logging.error(f"  鉴定点编号: {knowledge_point.get('编号', '未知')}")
            logging.error(f"  鉴定点名称: {knowledge_point.get('名称', '未知')}")
            sys.stdout.flush()
            return None
        
        requirements = self.LEVEL_REQUIREMENTS[level]

        # 构建prompt，一次性生成所有题目
        prompt = f"""
你是一个专业的考试命题专家。请根据以下知识点内容，一次性生成该鉴定点的全部考题。

【知识点信息】
鉴定点编号: {knowledge_point['编号']}
鉴定点名称: {knowledge_point['名称']}
知识点内容: {knowledge_point['内容']}

【特别注意】
⚠️ 知识点内容中可能包含规范书名和章节编号（如"《混凝土结构设计规范》9.2.4条"），这些仅是**参考来源**：
- 书名和章节号**只能出现在知识点内容中**，作为说明来源
- **题干和选项中必须完全删除这些章节编号**
- **必须从知识点内容中提取实质性的规定、数值、方法、要求等作为考点**
- **绝对不能考查"遵循哪一条规定"这种要求记住章节编号的题目**

📌 处理示例：
知识点内容："《混凝土结构设计规范》9.2.4条规定，应有不少于2根上部钢筋伸至外端，并向下弯折不小于12d"
- ❌ 错误出题："（   ）规定了钢筋应伸至外端" → 选项都是章节号
- ✅ 正确出题："悬臂梁中应有不少于（   ）根上部钢筋伸至外端" → 选项是"1/2/3/4"
- ✅ 正确出题："上部钢筋向下弯折不小于（   ）" → 选项是"10d/12d/15d/20d"

【出题要求】
1. 题目难度适中
2. 题目内容必须严格基于提供的知识点内容（提取实质内容，忽略章节编号）
3. ⚠️ 【核心要求】题目必须聚焦核心知识点
   - **识别核心知识点**: 仔细分析知识点内容，找出最核心、最关键的概念、原则、方法、定义等
   - **题目聚焦核心**: 题目的考查点必须是这些核心知识点，而不是句尾的修饰性内容
   - **避免形式主义**: 不要机械地将知识点内容的最后部分作为考查点
   - 📌 **示例说明**:
     * 如果知识点内容是"责任意识与担当精神是行为的基石与风骨"
     * ❌ 错误做法: 责任意识与担当精神是行为的（   ）[考查"基石与风骨"这个句尾内容]
     * ✅ 正确做法: （   ）是行为的基石与风骨 [考查核心"责任意识与担当精神"]
4. ⚠️ 【关键】正确答案必须在知识点原文中直接出现
   - 正确答案的表述必须在知识点原文中逐字逐句地出现
   - 不允许推导、不允许引申、不允许联想
   - 不能使用知识点中未提及的专业术语或概念作为答案
   - 不能基于外部知识或常识编造答案
   - 如果知识点中没有明确写出某个概念，就不能将其作为答案
   - 违反此规则的题目一律无效
4. 题干必须是陈述句，以句号结束，不能是疑问句
5. 题干中不能有换行符
6. 相同知识点可以用不同题目进行考查，但严禁题目重复。
7. ⚠️ 【禁止】题干中严禁出现"鉴定点"、"知识点"等指向不明的词汇
   - 不能出现"鉴定点中"、"知识点中"、"该知识点"等表述
   - 题干必须直接描述具体内容，不能引用"知识点"这个元概念
   - ❌ 错误示例：诚信正直的品格在知识点中被描述为（   ）
   - ✅ 正确示例：诚信正直的品格被描述为（   ）
8. ⚠️ 【禁止】题干中严禁出现"（   ）等"这种模糊表述
   - "（   ）等"暗示还有其他未列举的选项，会造成答案不唯一或有争议
   - 题目必须聚焦在确定性的核心概念上，不能让学生猜测"还有哪些"
   - ❌ 错误示例：行为锚定等级评价法针对（   ）等维度，为每个绩效等级锚定典型行为范例
   - ✅ 正确做法：考查确定性概念，如"（   ）是为每个绩效等级锚定典型行为范例的方法"
   - 📌 判断标准：如果题干包含"（   ）等"，这题就是违规的
9. ⚠️ 【禁止】题干和选项中严禁出现图表引用
   - 不能出现"图X.X"、"表X.X"、"附图"、"见图"等图表引用
   - 不能出现"如图所示"、"参见表格"等表述
   - 所有题目必须是纯文字描述，不依赖任何图表
10. ⚠️ 【禁止】严禁出"例如"题或"举例"题
   - ❌ 严禁题干中出现"例如（   ）"这种让学生选择例子的格式
   - ❌ 严禁让学生从几个例子中选择知识点给出的例子
   - ❌ 严禁考察"例如XX、YY"中的具体例子是哪个
   - ✅ 题目应该考察概念、规则、方法，而不是记住例子
   - ✅ 如果知识点只给了例子没有解释，应该考察例子背后的规律或特征
   - 📌 判断标准：如果题干包含"例如（   ）"或"如（   ）"，这题就是违规的
11. ⚠️ 【禁止】题干和选项中严禁出现章节编号或书籍目录信息
   - ❌ 题干中不能出现"5.2.4"、"第3章"、"9.2.7条"等章节编号
   - ❌ 题干中不能出现"根据5.2.4规定"、"按照第3章要求"等引用章节的表述
   - ❌ **选项中绝对不能出现章节编号**（如"9.2.7条"、"5.2.4"、"第3章"等）
   - ❌ **严禁考查"应该遵循哪一条规定"这类要求学生记住章节编号的题目**
   - 原因：
     * 章节信息来自参考书目录，会引导学生查找书籍
     * 考查章节编号本身毫无意义，是低质量出题
     * 应该考查知识点的实质内容，而不是它在书中的位置
   - ✅ 正确做法：
     * 从知识点内容中提取实质性的规定、要求、数值、方法等
     * 直接陈述知识点内容，不引用章节号
     * 考查学生对知识点本身的理解，而非对目录的记忆
   - 📌 示例：
     * ❌ 错误题干："根据9.2.4条规定，上柱插筋应包括在（   ）中"
     * ❌ 错误选项："A) 9.2.7条  B) 9.2.4条  C) 9.2.8条  D) 9.1.3条"
     * ✅ 正确题干："钢筋混凝土悬臂梁中，应有不少于（   ）根上部钢筋伸至悬臂梁外端"
     * ✅ 正确选项："A) 1  B) 2  C) 3  D) 4"
12. ⚠️ 【禁止】题干和选项中严禁出现具体图纸引用
   - ❌ 不能出现"如图一所示"、"图X所示"、"根据图纸"等引用具体图纸的表述
   - ❌ 不能出现"见附图"、"参考示意图"、"下图中"等需要配图的表述
   - ❌ 不能出现"图中的XX"、"图示结构"等依赖图纸理解的内容
   - 原因：考题中不附带详细的图纸，学生无法看到图纸内容
   - ✅ 正确做法：用文字完整描述所有信息，确保不看图也能理解和作答
   - 📌 示例：
     * ❌ 错误："如图一所示，柱变截面处的钢筋应如何配置？"
     * ✅ 正确："柱变截面处上柱插筋与下柱钢筋一同安装时，上柱插筋应包括在哪里？"

【知识覆盖要求】
必须确保题目覆盖以下两个层面：

一、基础且核心的知识点（宏观层面）：
   - 基本定义：该知识点的核心概念是什么
   - 分类体系：分为哪几类、有哪些类型
   - 组成框架：由哪些部分构成、包含哪些要素
   - 表达方式：如何表示、如何命名
   - 整体逻辑：遵循什么原则、有什么规律

二、具体内容细节（微观层面）：
   - 关键参数：具体的数值、尺寸、范围
   - 符号规则：使用什么符号、符号含义
   - 标注方法：如何书写、如何标注
   - 数值规定：具体规定的数值或比例
   - 条件限制：什么情况下适用、有何限制
   - 例外情况：特殊情况的处理
   - 典型示例：常见的应用实例

💡 出题时必须确保：
   - 既有对宏观知识点的考查，也有对微观知识点的考查
   - 题目应全面覆盖知识点的理论和实践两个维度

【题目数量】
- 单选题：{requirements['单选']}道（每题4个选项，有且只有一个正确答案）
- 判断题：{requirements['判断']}道（答案为"正确"或"错误"）
- 多选题：{requirements['多选']}道（每题5个选项，有两个或以上正确答案）

【单选题要求】
- 题干中要有且只有一个括号"（   ）"用于填空
- ⚠️ 括号位置要有多样性：在同一鉴定点的所有单选题和多选题中，括号不能都出现在题干尾部
  * 至少要有部分题目的括号出现在题干的开头、中间等不同位置
  * 这样可以从多角度考查知识点，避免出题模式单一
- 必须提供4个选项(A、B、C、D)
- 有且只有一个正确答案
- ⚠️ 4个选项的内容必须各不相同，不能有重复
- 其他三个选项应该具有一定的干扰性，不能明显错误
- 选项内容简洁，不换行

【判断题要求】
- 题干是一个陈述句，以句号结束
- 不需要括号
- 不需要选项
- 题目的陈述可以是正确的(答案为"正确")，也可以是错误的(答案为"错误")
- 如果是错误的陈述，错误点应该具有一定的迷惑性

【多选题要求】
- 题干中要有且只有一个括号"（   ）"用于填空
- ⚠️ 括号位置要有多样性：在同一鉴定点的所有单选题和多选题中，括号不能都出现在题干尾部
  * 至少要有部分题目的括号出现在题干的开头、中间等不同位置
  * 这样可以从多角度考查知识点，避免出题模式单一
- 必须提供5个选项(A、B、C、D、E)
- 正确答案有两个或两个以上
- ⚠️ 5个选项的内容必须各不相同，不能有重复
- 其他选项应该具有一定的干扰性
- 选项内容简洁，不换行
- 答案字段中多个选项字母之间不要用间隔符号（如：ABC、BDE）

【输出格式】
请以JSON格式输出，包含一个"题目列表"数组，每个题目包含以下字段：
{{
    "题目列表": [
        {{
            "题目类型": "单选",
            "题干": "题目内容（   ）。",
            "选项A": "选项A内容",
            "选项B": "选项B内容",
            "选项C": "选项C内容",
            "选项D": "选项D内容",
            "答案": "B"
        }},
        {{
            "题目类型": "判断",
            "题干": "题目陈述内容。",
            "答案": "正确"
        }},
        {{
            "题目类型": "多选",
            "题干": "题目内容（   ）。",
            "选项A": "选项A内容",
            "选项B": "选项B内容",
            "选项C": "选项C内容",
            "选项D": "选项D内容",
            "选项E": "选项E内容",
            "答案": "ABC"
        }}
    ]
}}

注意：
1. 确保题目之间没有重复，同一题型不能考查同一个知识点，但不同题型可以考查同一个知识点。
2. 题目顺序：先单选题，再判断题，最后多选题
3. 严格按照上述JSON格式输出
"""

        try:
            logging.info(f"  正在调用AI生成题目...")
            sys.stdout.flush()  # 立即刷新输出

            completion = self.client.chat.completions.create(
                model="qwen3-max",
                messages=[
                    {
                        "role": "system",
                        "content": "你是一个专业的考试命题专家。请严格按照要求生成考题，确保题目质量高、没有重复，以JSON格式返回。"
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                response_format={"type": "json_object"}
            )

            logging.info(f"  AI响应完成，正在解析结果...")
            sys.stdout.flush()  # 立即刷新输出

            json_string = completion.choices[0].message.content
            
            # 记录原始响应（用于调试）
            logging.debug(f"  AI原始响应: {json_string[:500]}...")
            
            response_data = json.loads(json_string)

            # 提取题目列表
            questions = response_data.get("题目列表", [])
            
            # 检查是否成功生成题目
            if not questions:
                logging.warning(f"  ⚠️  AI返回了空的题目列表")
                logging.info(f"  完整响应: {json_string}")
                return None

            # 为每个题目添加鉴定点编号
            for q in questions:
                q["鉴定点编号"] = knowledge_point["编号"]

            return questions

        except json.JSONDecodeError as e:
            logging.error(f"  ❌ JSON解析失败: {e}")
            logging.error(f"  AI响应内容: {json_string[:1000] if 'json_string' in locals() else '未获取到响应'}")
            sys.stdout.flush()
            return None
        except KeyError as e:
            logging.error(f"  ❌ 响应格式错误，缺少必要字段: {e}")
            logging.error(f"  响应内容: {response_data if 'response_data' in locals() else '未解析'}")
            sys.stdout.flush()
            return None
        except Exception as e:
            logging.error(f"  ❌ 生成题目时出错: {type(e).__name__}: {e}")
            sys.stdout.flush()
            return None

    def evaluate_questions(self, knowledge_point: Dict[str, str],
                          questions: List[Dict]) -> Dict:
        """
        评估生成的题目质量

        Args:
            knowledge_point: 鉴定点信息
            questions: 生成的题目列表

        Returns:
            评估结果，包含问题和建议
        """
        # 构建评估prompt
        questions_text = json.dumps(questions, ensure_ascii=False, indent=2)

        prompt = f"""
你是一个专业的考试质量评估专家。请对以下生成的考题进行全面评估。

【知识点信息】
鉴定点编号: {knowledge_point['编号']}
鉴定点名称: {knowledge_point['名称']}
知识点内容: {knowledge_point['内容']}

【生成的题目】
{questions_text}

【评估维度】
请从以下维度评估题目质量：

1. **有效性检查和重复性检查**（⚠️ 最重要）
   - ⚠️ 【关键】题目的正确答案是否在知识点原文中逐字逐句地出现？
     * 正确答案必须在知识点原文中直接出现，逐字可查
     * 不允许推导、不允许引申、不允许联想得出答案
     * 不能使用知识点中未提及的专业术语或概念
     * 不能基于外部知识编造答案
     * 这是最重要的检查项，违反此项必须标记为严重问题
   - ⚠️ 【禁止】题干和选项中是否出现了图表引用？
     * 不能出现"图X.X"、"表X.X"、"附图"、"见图"等
     * 不能出现"如图所示"、"参见表格"等表述
     * 违反此项必须标记为严重问题
   - ⚠️ 【禁止】是否出现了"例如"题或"举例"题？
     * ❌ 检查题干中是否包含"例如（   ）"或"如（   ）"格式
     * ❌ 检查是否让学生从几个例子中选择知识点给出的例子
     * ❌ 检查是否考察"例如XX、YY"中的具体例子是哪个
     * 📌 特别注意：如果题干是"XX例如（   ）"，答案是知识点中的例子，这就是违规
     * 违反此项必须标记为严重问题
   - ⚠️ 【禁止】题干和选项中是否出现了章节编号？【重要】
     * ❌ 题干中是否包含"5.2.4"、"第3章"、"9.2.7条"等章节编号
     * ❌ 题干中是否有"根据5.2.4规定"、"按照第3章要求"等表述
     * ❌ **选项中是否出现章节编号**（如选项为"A) 9.2.7条"、"B) 5.2.4"等）
     * ❌ **是否在考查"应遵循哪一条规定"这类要求记住章节编号的题目**
     * 特别检查：如果选项都是"X.X.X条"格式，这是严重的低质量出题
     * 原因：
       - 章节信息会引导学生查找书籍而非考查知识点本身
       - 考查章节编号本身毫无意义，是最低质量的出题
       - 应该考查知识点的实质内容（数值、方法、要求等）
     * 违反此项必须标记为**严重问题**
   - ⚠️ 【禁止】题干和选项中是否出现了具体图纸引用？
     * ❌ 检查是否包含"如图一所示"、"图X所示"、"根据图纸"等表述
     * ❌ 检查是否有"见附图"、"参考示意图"、"下图中"等需要配图的内容
     * ❌ 检查是否有"图中的XX"、"图示结构"等依赖图纸的描述
     * 原因：考题不附带图纸，学生无法看到图纸内容
     * 违反此项必须标记为严重问题
   - 是否有同一题型下题目内容重复或高度相似？
   - 是否有题目内容与知识点内容不符？

2. **核心知识点考察**（⭐ 重要）
   - ⚠️ 题目是否聚焦到核心知识点上？
     * 检查题目考查的是否是知识点中最核心、最关键的概念、原则、方法、定义等
     * 检查题目是否避免了机械地考查句尾的修饰性内容
     * 📌 判断标准：如果知识点是"XX是YY"，题目应考查核心的"XX"，而不是描述性的"YY"
     * 例如："责任意识与担当精神是行为的基石与风骨"，应考查"责任意识与担当精神"而非"基石与风骨"
   - 题目是否准确考察了核心知识点？
   - 题目是否涵盖了知识点的主要方面？

3. **知识覆盖度评估**（⭐ 重要）
   请检查题目是否全面覆盖以下两个层面：

   宏观层面（基础核心知识）：
   - 是否考察了基本定义、核心概念？
   - 是否考察了分类体系、类型划分？
   - 是否考察了组成框架、构成要素？
   - 是否考察了表达方式、命名规则？
   - 是否考察了整体逻辑、遵循原则？

   微观层面（具体内容细节）：
   - 是否考察了关键参数、具体数值？
   - 是否考察了符号规则、标注方法？
   - 是否考察了条件限制、适用情况？
   - 是否考察了例外情况、特殊处理？
   - 是否考察了典型示例、实际应用？

   💡 评估要点：
   - 题目宏观考查与微观考查都要有。
   - 如果只偏重某一层面，应指出缺失的方面

4. **选项准确度**（适用于单选题和多选题）
   - ⚠️ 【关键】正确答案是否在知识点原文中逐字逐句地出现？
     * 答案必须在知识点原文中直接出现，逐字可查
     * 不允许推导、引申或联想得出答案
     * 答案中的专业术语必须在知识点中明确出现过
     * 不能使用知识点未提及的概念作为答案
   - 干扰选项是否合理且具有一定迷惑性？
   - 选项是否有明显错误或不合理之处？
   - 多选题的正确答案数量是否合理（至少2个）？

5. **格式规范性**
   - 题干是否为陈述句，以句号结束？
   - 单选/多选题是否有且只有一个括号（   ）？
   - ⚠️ 括号位置多样性：同一鉴定点的所有单选题和多选题中，括号是否都出现在题干尾部？
     * 如果所有选择题的括号都在尾部，这是严重问题
     * 应该有部分题目的括号在开头或中间位置
   - 单选题是否有4个选项？多选题是否有5个选项？
   - ⚠️ 单选题的4个选项内容是否各不相同？多选题的5个选项内容是否各不相同？
   - 答案格式是否正确（单选：A，多选：ABC，判断：正确/错误）？
   - ⚠️ 题干中是否出现了"鉴定点"、"知识点"等指向不明的词汇？
     * 如出现"知识点中"、"该知识点"、"鉴定点描述"等，必须标记为严重问题
   - ⚠️ 题干中是否出现了"（   ）等"这种模糊表述？
     * 这种表述暗示还有其他未列举选项，会造成答案不唯一或有争议
     * 必须标记为严重问题
   - ⚠️ 题干和选项中是否出现了图表引用（如"图X.X"、"表X.X"、"如图所示"等）？
   - ⚠️ 是否出现了"例如"题？检查题干中是否包含"例如（   ）"或"如（   ）"格式。
   - ⚠️ 题干和选项中是否出现了章节编号？
     * 检查是否有"5.2.4"、"第3章"、"3.1.2"等章节编号
     * 如有，必须标记为严重问题
   - ⚠️ 题干和选项中是否出现了具体图纸引用？
     * 检查是否有"如图一所示"、"图X所示"、"见附图"、"图中的XX"等表述
     * 如有，必须标记为严重问题

6. **语言表达**
   - 题干和选项的表达是否清晰准确？
   - 是否存在歧义或语病？

【输出格式】
请以JSON格式输出评估结果：
{{
    "总体评价": "优秀/良好/需要改进",
    "是否通过": true/false,
    "问题列表": [
        {{
            "题目序号": 1,
            "问题类型": "重复性/核心知识点/知识覆盖度/选项准确度/格式规范/语言表达",
            "问题描述": "具体问题说明",
            "严重程度": "严重/一般/轻微"
        }}
    ],
    "修改建议": [
        {{
            "题目序号": 1,
            "建议内容": "具体的修改建议"
        }}
    ]
}}

注意：
1. 如果没有问题，问题列表为空数组，is_passed为true
2. 严重问题必须修改，一般问题建议修改，轻微问题可以忽略
3. 有严重问题或多个一般问题时，is_passed应为false
"""

        try:
            logging.info(f"  正在调用AI评估题目...")
            sys.stdout.flush()  # 立即刷新输出

            completion = self.client.chat.completions.create(
                model="qwen3-max",
                messages=[
                    {
                        "role": "system",
                        "content": "你是一个专业的考试质量评估专家。请客观、严格地评估题目质量，以JSON格式返回评估结果。"
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                response_format={"type": "json_object"}
            )

            logging.info(f"  AI评估完成，正在解析结果...")
            sys.stdout.flush()  # 立即刷新输出

            json_string = completion.choices[0].message.content
            evaluation = json.loads(json_string)

            return evaluation

        except Exception as e:
            logging.error(f"评估题目时出错: {e}")
            return {"是否通过": True, "问题列表": [], "修改建议": []}

    def fix_questions(self, knowledge_point: Dict[str, str],
                     questions: List[Dict],
                     evaluation: Dict) -> Optional[List[Dict]]:
        """
        根据评估结果修正题目

        Args:
            knowledge_point: 鉴定点信息
            questions: 原始题目列表
            evaluation: 评估结果

        Returns:
            修正后的题目列表
        """
        questions_text = json.dumps(questions, ensure_ascii=False, indent=2)
        problems_text = json.dumps(evaluation.get("问题列表", []), ensure_ascii=False, indent=2)
        suggestions_text = json.dumps(evaluation.get("修改建议", []), ensure_ascii=False, indent=2)

        prompt = f"""
你是一个专业的考试命题专家。请根据评估反馈，修正以下题目中存在的问题。

【知识点信息】
鉴定点编号: {knowledge_point['编号']}
鉴定点名称: {knowledge_point['名称']}
知识点内容: {knowledge_point['内容']}

【原始题目】
{questions_text}

【发现的问题】
{problems_text}

【修改建议】
{suggestions_text}

【修正要求】
1. 针对每个问题进行修正
2. 保持题目总数不变
3. 确保修正后的题目符合所有格式要求
4. 题目内容必须基于知识点内容
5. 避免题目重复

【输出格式】
请以JSON格式输出修正后的完整题目列表：
{{
    "题目列表": [
        {{
            "题目类型": "单选/判断/多选",
            "题干": "...",
            "选项A": "...",
            "答案": "..."
        }}
    ]
}}
"""

        try:
            logging.info(f"  正在调用AI修正题目...")
            sys.stdout.flush()  # 立即刷新输出

            completion = self.client.chat.completions.create(
                model="qwen3-max",
                messages=[
                    {
                        "role": "system",
                        "content": "你是一个专业的考试命题专家。请根据反馈认真修正题目，确保题目质量，以JSON格式返回。"
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                response_format={"type": "json_object"}
            )

            logging.info(f"  AI修正完成，正在解析结果...")
            sys.stdout.flush()  # 立即刷新输出

            json_string = completion.choices[0].message.content
            response_data = json.loads(json_string)

            # 提取题目列表
            fixed_questions = response_data.get("题目列表", [])

            # 为每个题目添加鉴定点编号
            for q in fixed_questions:
                q["鉴定点编号"] = knowledge_point["编号"]

            return fixed_questions

        except Exception as e:
            logging.error(f"修正题目时出错: {e}")
            return None

    def generate_questions_for_point(self, knowledge_point: Dict[str, str],
                                    level: str) -> List[Dict]:
        """
        为单个鉴定点生成所有题目（包含质量评估和修正机制）
        要求：必须达到"优秀"评级才能保存

        Args:
            knowledge_point: 鉴定点信息
            level: 等级（仅用于确定题目数量：一级/二级/三级=7题，四级/五级=6题）

        Returns:
            题目列表
        """
        logging.info(f"\n正在为鉴定点 {knowledge_point['编号']} 生成题目...")
        sys.stdout.flush()

        max_iterations = 5  # 增加最大迭代次数，确保有足够机会达到优秀
        iteration = 0

        while iteration < max_iterations:
            iteration += 1
            logging.info(f"\n  第 {iteration} 轮生成...")
            sys.stdout.flush()

            # 生成题目
            questions = self.generate_all_questions_at_once(knowledge_point, level)

            if not questions:
                logging.warning(f"  ❌ 生成失败")
                sys.stdout.flush()
                continue

            logging.info(f"  ✓ 成功生成 {len(questions)} 道题目")
            sys.stdout.flush()

            # 评估题目质量
            logging.info(f"  正在评估题目质量...")
            sys.stdout.flush()
            evaluation = self.evaluate_questions(knowledge_point, questions)

            is_passed = evaluation.get("是否通过", False)
            problems = evaluation.get("问题列表", [])
            overall = evaluation.get("总体评价", "未知")

            logging.info(f"  评估结果: {overall}")
            sys.stdout.flush()

            # 只接受"优秀"评级
            if overall == "优秀" and is_passed:
                logging.info(f"  ✓ 题目质量优秀，通过评估！")
                sys.stdout.flush()
                return questions
            elif overall == "良好":
                logging.warning(f"  ⚠ 题目质量为良好，需要优化至优秀")
                logging.warning(f"  ⚠ 发现 {len(problems)} 个问题，需要改进")
                sys.stdout.flush()
            else:
                logging.warning(f"  ⚠ 题目质量需要改进，发现 {len(problems)} 个问题")
                sys.stdout.flush()

            # 显示问题
            for i, problem in enumerate(problems[:3], 1):  # 只显示前3个问题
                logging.info(f"    - 题目{problem.get('题目序号', '?')}: {problem.get('问题描述', '')[:50]}...")
            sys.stdout.flush()

            if iteration < max_iterations:
                logging.info(f"  正在优化题目...")
                sys.stdout.flush()
                fixed_questions = self.fix_questions(knowledge_point, questions, evaluation)

                if fixed_questions:
                    questions = fixed_questions
                    logging.info(f"  ✓ 题目已优化，进入下一轮评估")
                    sys.stdout.flush()
                else:
                    logging.warning(f"  ❌ 优化失败，重新生成")
                    sys.stdout.flush()
                    continue
            else:
                logging.warning(f"  ⚠ 已达到最大迭代次数（{max_iterations}轮）")
                if overall == "良好":
                    logging.warning(f"  ⚠ 当前评级为良好，建议手动检查")
                sys.stdout.flush()
                return questions

        logging.error(f"  ❌ 未能生成优秀级别的题目")
        sys.stdout.flush()
        return []

    def save_questions_to_xls(self, questions: List[Dict],
                              output_filepath: str):
        """
        将题目保存到XLS文件

        Args:
            questions: 题目列表
            output_filepath: 输出文件路径
        """
        # 创建工作簿
        wb = xlwt.Workbook(encoding='utf-8')
        ws = wb.add_sheet('题目')

        # 设置宋体样式
        style = xlwt.XFStyle()
        font = xlwt.Font()
        font.name = '宋体'
        font.height = 220  # 11号字体 (20 * 11)
        style.font = font

        # 写入表头
        headers = ["鉴定点代码", "题目类型代码", "试题(题干)",
                  "选项A", "选项B", "选项C", "选项D", "选项E",
                  "答案", "难度代码", "一致性代码"]
        for col, header in enumerate(headers):
            ws.write(0, col, header, style)

        # 写入题目
        for row_idx, q in enumerate(questions, start=1):
            # 确定题目类型代码
            type_code_map = {"单选": "B", "判断": "C", "多选": "D"}
            type_code = type_code_map.get(q.get("题目类型", ""), "B")

            # 写入数据
            data = [
                q.get("鉴定点编号", ""),
                type_code,
                q.get("题干", ""),
                q.get("选项A", ""),
                q.get("选项B", ""),
                q.get("选项C", ""),
                q.get("选项D", ""),
                q.get("选项E", ""),
                q.get("答案", ""),
                3,  # 难度代码
                5   # 一致性代码
            ]
            for col_idx, value in enumerate(data):
                ws.write(row_idx, col_idx, value, style)

        wb.save(output_filepath)
        logging.info(f"  ✓ XLS文件已保存到: {output_filepath}")

    def save_questions_to_docx(self, questions: List[Dict],
                               output_filepath: str):
        """
        将题目保存到Word文档

        Args:
            questions: 题目列表
            output_filepath: 输出文件路径
        """
        doc = Document()

        # 设置默认字体为宋体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(11)

        # 确定题目类型代码映射
        type_code_map = {"单选": "B", "判断": "C", "多选": "D"}

        # 逐题写入
        for idx, q in enumerate(questions):
            # 获取题目信息
            code = q.get("鉴定点编号", "")
            type_code = type_code_map.get(q.get("题目类型", ""), "B")
            question_text = q.get("题干", "")
            answer = q.get("答案", "")

            # 第1行: 鉴定点编号  题型  难度代码  一致性代码
            p1 = doc.add_paragraph(f"{code}  {type_code}  3  5")
            p1.style.font.name = '宋体'
            p1.style.font.size = Pt(11)

            # 第2行: {A} 题干
            p2 = doc.add_paragraph(f"{{A}}{question_text}")
            p2.style.font.name = '宋体'
            p2.style.font.size = Pt(11)

            # 选项行 (单选和多选题)
            if type_code in ["B", "D"]:  # 单选或多选
                for option_key in ["A", "B", "C", "D", "E"]:
                    option_value = q.get(f"选项{option_key}", "")
                    if option_value:  # 只写入非空选项
                        p_opt = doc.add_paragraph(f"（{option_key}）{option_value}")
                        p_opt.style.font.name = '宋体'
                        p_opt.style.font.size = Pt(11)

            # 答案行: {B} 答案
            p_ans = doc.add_paragraph(f"{{B}}{answer}")
            p_ans.style.font.name = '宋体'
            p_ans.style.font.size = Pt(11)

            # 题目间空一行 (除了最后一题)
            if idx < len(questions) - 1:
                doc.add_paragraph()

        doc.save(output_filepath)
        logging.info(f"  ✓ Word文档已保存到: {output_filepath}")

    def get_existing_question_codes(self, output_dir: str) -> set:
        """
        获取已存在的题目文件对应的鉴定点编号
        
        Args:
            output_dir: 输出目录路径（包含xlsx文件名的子目录）
            
        Returns:
            已存在题目的鉴定点编号集合
        """
        existing_codes = set()
        
        if not os.path.exists(output_dir):
            return existing_codes
            
        for filename in os.listdir(output_dir):
            if filename.startswith("考题") and filename.endswith(".xls"):
                # 从文件名中提取鉴定点编号，格式：考题A-B-A-001.xls
                code = filename[2:-4]  # 去掉"考题"前缀和".xls"后缀
                existing_codes.add(code)
                
        return existing_codes

    def resolve_input_path(self, input_path: str) -> str:
        """
        解析输入文件路径，如果只提供文件名，则在resources目录中查找

        Args:
            input_path: 输入路径或文件名

        Returns:
            完整的文件路径
        """
        # 如果是绝对路径或相对路径（包含路径分隔符），直接使用
        if os.path.sep in input_path or os.path.isabs(input_path):
            return input_path
        
        # 如果只是文件名，在resources目录中查找
        resources_dir = os.path.join(os.path.dirname(__file__), "resources")
        resources_path = os.path.join(resources_dir, input_path)
        
        if os.path.exists(resources_path):
            return resources_path
        
        # 如果resources目录中没有，返回原路径
        return input_path

    def process_file(self, input_filepath: str, output_dir: str = None):
        """
        处理单个输入文件,生成考题

        Args:
            input_filepath: 输入xlsx文件路径
            output_dir: 输出目录,如果为None则默认使用questions目录
        """
        # 解析输入文件路径
        resolved_input_path = self.resolve_input_path(input_filepath)
        
        # 检查文件是否存在
        if not os.path.exists(resolved_input_path):
            logging.error(f"❌ 错误: 文件不存在: {resolved_input_path}")
            sys.stdout.flush()
            return
        
        # 提取文件名
        filename = os.path.basename(resolved_input_path)
        logging.info(f"\n{'='*60}")
        logging.info(f"处理文件: {filename}")
        logging.info(f"文件路径: {resolved_input_path}")
        logging.info(f"{'='*60}")
        sys.stdout.flush()

        # 提取等级信息
        level = self.extract_level_from_filename(filename)
        if not level:
            logging.error(f"❌ 错误: 无法从文件名 '{filename}' 中提取等级信息")
            logging.error(f"   文件名必须包含: 一级、二级、三级、四级或五级")
            sys.stdout.flush()
            return

        logging.info(f"✓ 检测到等级: {level}")
        logging.info(f"✓ 出题要求: {self.LEVEL_REQUIREMENTS[level]}")
        sys.stdout.flush()

        # 读取鉴定点
        logging.info(f"正在读取鉴定点...")
        sys.stdout.flush()
        knowledge_points = self.read_knowledge_points(resolved_input_path)
        logging.info(f"✓ 读取到 {len(knowledge_points)} 个鉴定点")
        sys.stdout.flush()

        # 确定输出目录 - 在questions下创建以xlsx文件名命名的子目录
        if output_dir is None:
            # 默认保存到questions目录
            base_output_dir = os.path.join(os.path.dirname(__file__), "questions")
            # 从文件名中提取基础名称（去掉.xlsx后缀）
            xlsx_basename = os.path.splitext(filename)[0]
            # 创建子目录：questions/五级5001-5010/
            output_dir = os.path.join(base_output_dir, xlsx_basename)

        # 确保输出目录存在
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            logging.info(f"✓ 创建输出目录: {output_dir}")

        logging.info(f"✓ 输出目录: {output_dir}")
        sys.stdout.flush()

        # 获取已存在的题目文件（从当前xlsx对应的子目录中）
        existing_codes = self.get_existing_question_codes(output_dir)
        if existing_codes:
            logging.info(f"✓ 发现已存在 {len(existing_codes)} 个题目文件")
            logging.info(f"  已存在的鉴定点: {', '.join(sorted(existing_codes))}")
        else:
            logging.info(f"✓ 未发现已存在的题目文件，将生成所有题目")
        sys.stdout.flush()

        # 筛选需要生成题目的鉴定点（增量生成）
        new_knowledge_points = [kp for kp in knowledge_points if kp['编号'] not in existing_codes]
        
        if not new_knowledge_points:
            logging.info(f"✓ 所有鉴定点的题目都已存在，无需生成新题目")
            sys.stdout.flush()
            return
        
        logging.info(f"✓ 需要生成题目的鉴定点: {len(new_knowledge_points)} 个")
        sys.stdout.flush()

        # 为需要生成的鉴定点生成题目并保存
        for i, kp in enumerate(new_knowledge_points, 1):
            logging.info(f"\n{'─'*60}")
            logging.info(f"鉴定点: {kp['编号']} - {kp['名称']}")
            sys.stdout.flush()

            # 生成题目
            questions = self.generate_questions_for_point(kp, level)

            if not questions:
                logging.warning(f"  ⚠ 警告: 该鉴定点未生成任何题目")
                sys.stdout.flush()
                continue

            logging.info(f"  ✓ 成功生成 {len(questions)} 道题目")
            sys.stdout.flush()

            # 保存为XLS格式（保持原有格式）
            xls_filename = f"考题{kp['编号']}.xls"
            xls_filepath = os.path.join(output_dir, xls_filename)
            self.save_questions_to_xls(questions, xls_filepath)

            # 保存为Word文档
            docx_filename = f"{kp['编号']}.docx"
            docx_filepath = os.path.join(output_dir, docx_filename)
            self.save_questions_to_docx(questions, docx_filepath)


def process_all_resources():
    """处理resources目录中的所有xlsx文件"""
    # 初始化日志系统
    log_filepath = setup_logging()
    logging.info("="*60)
    logging.info("考题生成系统启动 - 批量增量处理模式")
    logging.info(f"日志文件: {log_filepath}")
    logging.info("="*60)
    sys.stdout.flush()

    # 获取resources目录路径
    script_dir = os.path.dirname(__file__)
    resources_dir = os.path.join(script_dir, "resources")
    
    if not os.path.exists(resources_dir):
        logging.error(f"❌ 错误: resources目录不存在: {resources_dir}")
        sys.stdout.flush()
        return

    # 获取所有xlsx文件
    xlsx_files = [f for f in os.listdir(resources_dir) if f.endswith('.xlsx')]
    
    if not xlsx_files:
        logging.info("✓ resources目录中没有找到xlsx文件")
        sys.stdout.flush()
        return

    logging.info(f"✓ 发现 {len(xlsx_files)} 个教材文件:")
    for f in sorted(xlsx_files):
        logging.info(f"  - {f}")
    sys.stdout.flush()

    # 创建生成器
    try:
        generator = QuestionGenerator()
    except ValueError as e:
        logging.error(f"❌ 初始化失败: {e}")
        logging.error("请确保在.env文件中设置了DASHSCOPE_API_KEY")
        sys.stdout.flush()
        return

    # 处理每个文件
    total_processed = 0
    for xlsx_file in sorted(xlsx_files):
        try:
            generator.process_file(xlsx_file)
            total_processed += 1
        except Exception as e:
            logging.error(f"❌ 处理文件 {xlsx_file} 时出错: {e}")
            sys.stdout.flush()
            continue

    logging.info(f"\n{'='*60}")
    logging.info(f"批量处理完成! 成功处理 {total_processed}/{len(xlsx_files)} 个文件")
    logging.info(f"日志已保存到: {log_filepath}")
    logging.info(f"{'='*60}\n")

def main():
    """主函数"""
    import sys

    # 检查命令行参数
    if len(sys.argv) < 2:
        # 没有参数时，批量处理resources目录中的所有文件
        process_all_resources()
        return

    # 有参数时，处理指定文件
    # 初始化日志系统
    log_filepath = setup_logging()
    logging.info("="*60)
    logging.info("考题生成系统启动 - 单文件处理模式")
    logging.info(f"日志文件: {log_filepath}")
    logging.info("="*60)
    sys.stdout.flush()

    input_file = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None

    # 创建生成器并处理文件
    try:
        generator = QuestionGenerator()
        generator.process_file(input_file, output_dir)
    except ValueError as e:
        logging.error(f"❌ 初始化失败: {e}")
        logging.error("请确保在.env文件中设置了DASHSCOPE_API_KEY")
        sys.stdout.flush()
        return

    logging.info(f"\n{'='*60}")
    logging.info("处理完成!")
    logging.info(f"日志已保存到: {log_filepath}")
    logging.info(f"{'='*60}\n")


if __name__ == "__main__":
    main()
